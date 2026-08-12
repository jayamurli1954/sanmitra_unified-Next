#!/usr/bin/env python3
"""Insert screenshots into docs/MitraBooks_User_Manual.docx at the right sections.

Drop PNGs into docs/User_manual_screenshots/, list them in MAPPING below (anchor heading ->
image file + caption), then run this script. Each image is placed immediately after
its section heading, scaled to page width, with an italic caption beneath it.

Idempotent: before inserting for an anchor, any existing figure (image paragraph(s)
+ "Figure ..." caption) directly under that heading is removed and re-inserted, so
re-running after replacing a PNG just updates it — no duplicates.

Only entries whose image file actually exists are inserted, so you can add
screenshots incrementally.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MANUAL = ROOT / "docs" / "MitraBooks_User_Manual.docx"
ASSET_COPY = ROOT / "frontend" / "assets" / "mitrabooks" / "MitraBooks_User_Manual.docx"
IMG_DIR = ROOT / "docs" / "User_manual_screenshots"

IMAGE_WIDTH = Inches(6.0)   # fits US-Letter / A4 with 1" margins

# (section-number token, image filename in docs/User_manual_screenshots/, caption).
# The token is the leading number of a heading: "3." for the H1 "3. Dashboard",
# "4.1" for "4.1 Adding a Party", etc. Matching by token is dash-proof. Drop a PNG
# with the given name into docs/User_manual_screenshots/ for any row you want illustrated;
# rows with no file are skipped, so this whole list can be filled incrementally.
MAPPING = [
    # Core walkthrough
    ("2.", "2.png", "Figure 2 — Getting Started"),
    ("2.1", "02-1-login.png", "Figure 2.1 — Logging in"),
    ("2.2", "02-2-workspace-layout.png", "Figure 2.2 — Workspace layout"),
    ("2.3", "2.3 Navigation Groups.png", "Figure 2.3 — Navigation groups (sidebar)"),
    ("3.", "03-dashboard.png", "Figure 3 — Dashboard workspace"),
    ("4.1", "04-1-add-party.png", "Figure 4.1 — Adding a party"),
    # Sales / Purchases
    ("5.1", "05-1-create-invoice.png", "Figure 5.1 — Creating a sales invoice"),
    ("5.3", "05-3-invoice-list.png", "Figure 5.3 — Invoice list"),
    ("5.4", "05-4-einvoice.png", "Figure 5.4 — e-Invoice (IRN)"),
    ("6.1", "06-1-credit-note.png", "Figure 6.1 — Creating a credit note"),
    ("7.1", "07-1-vendor-bill.png", "Figure 7.1 — Creating a vendor bill"),
    ("7.2", "07-2-debit-note.png", "Figure 7.2 — Debit notes"),
    # Payments / Banking
    ("8.1", "08-1-payment-allocation.png", "Figure 8.1 — Allocating a payment"),
    ("9.1", "09-1-bank-upload.png", "Figure 9.1 — Uploading a bank statement"),
    ("9.3", "09-3-brs.png", "Figure 9.3 — Bank Reconciliation Statement"),
    # Compliance
    ("10.1", "10-1-gstr3b.png", "Figure 10.1 — GSTR-3B"),
    ("10.2", "10-2-gstr1.png", "Figure 10.2 — GSTR-1"),
    ("10.3", "10-3-gstr2b.png", "Figure 10.3 — GSTR-2B / ITC reconciliation"),
    ("11.3", "11-3-tds-register.png", "Figure 11.3 — TDS / TCS register"),
    # Accounting / Reports
    ("12.1", "12-1-coa.png", "Figure 12.1 — Core Ledger (Chart of Accounts)"),
    ("12.2", "12-2-journal-post.png", "Figure 12.2 — Manual journal post"),
    ("12.3", "12-3-audit.png", "Figure 12.3 — Audit trails"),
    ("13.1", "13-1-trial-balance.png", "Figure 13.1 — Trial Balance"),
    ("13.2", "13-2-pnl.png", "Figure 13.2 — Profit & Loss"),
    ("13.3", "13-3-balance-sheet.png", "Figure 13.3 — Balance Sheet"),
    ("13.6", "13-6-aging.png", "Figure 13.6 — AR / AP Aging"),
    ("14.1", "14-1-statement.png", "Figure 14.1 — Customer / vendor statement"),
    # Period / Opening / Assets / Dimensions / Inventory / Settings / CA
    ("15.", "15-period-locks.png", "Figure 15 — Period locks"),
    ("16.1", "16-1-opening-balances.png", "Figure 16.1 — Uploading opening balances"),
    ("17.1", "17-1-asset-register.png", "Figure 17.1 — Asset register"),
    ("18.1", "18-1-dimensions.png", "Figure 18.1 — Setting up dimensions"),
    ("19.1", "19-1-item-master.png", "Figure 19.1 — Item master"),
    ("19.2", "19-2-stock-register.png", "Figure 19.2 — Stock register"),
    ("20.1", "20-1-settings.png", "Figure 20.1 — Invoice settings"),
    ("21.1", "21-1-ca-invite.png", "Figure 21.1 — Inviting a CA"),
    # HR & Payroll
    ("22.", "22-hr-workspace.png", "Figure 22 — HR & Payroll workspace"),
    ("22.3", "22-3-add-employee.png", "Figure 22.3 — Add Employee form"),
    ("22.5", "22-5-run-payroll.png", "Figure 22.5 — Payroll run"),
    ("22.6", "22-6-leave.png", "Figure 22.6 — Leave management"),
    ("22.8", "22-8-fnf.png", "Figure 22.8 — Full & Final settlement"),
    # Manufacturing & Cost Centres
    ("23.", "23-mfg-workspace.png", "Figure 23 — Manufacturing & Cost Centres workspace"),
    ("23.2", "23-2-cost-centres.png", "Figure 23.2 — Cost Centres tab"),
    ("23.4", "23-4-budgets.png", "Figure 23.4 — Budgets and budget-vs-actual"),
    ("23.5", "23-5-pl.png", "Figure 23.5 — Cost-Centre P&L report"),
    ("23.6", "23-6-bom.png", "Figure 23.6 — Creating a BOM"),
    ("23.7", "23-7-work-orders.png", "Figure 23.7 — Work Orders and completion"),
]


def _is_figure_para(p) -> bool:
    """A figure is an image paragraph (has a drawing) or our caption paragraph."""
    if p._p.findall(".//" + qn("w:drawing")):
        return True
    txt = p.text.strip()
    return txt.startswith("Figure ")


def _clear_existing_figures(doc, anchor):
    """Remove image/caption paragraphs directly following the anchor heading."""
    nxt = anchor._p.getnext()
    body = anchor._p.getparent()
    from docx.text.paragraph import Paragraph
    while nxt is not None and nxt.tag == qn("w:p"):
        para = Paragraph(nxt, anchor._parent)
        if _is_figure_para(para):
            after = nxt.getnext()
            body.remove(nxt)
            nxt = after
        else:
            break


def _insert_after(anchor, new_p):
    anchor._p.addnext(new_p._p)


def insert_figure(doc, anchor_token: str, image_path: Path, caption: str) -> bool:
    anchor = None
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text or not (p.style and p.style.name in ("Heading 1", "Heading 2")):
            continue
        # Match by the leading number token (e.g. "3." or "13.2") — dash-proof.
        if text.split(" ", 1)[0] == anchor_token:
            anchor = p
            break
    if anchor is None:
        print(f"  SKIP (anchor not found): {anchor_token!r}")
        return False

    _clear_existing_figures(doc, anchor)

    # Caption first, then image — both inserted right after the heading, so the
    # image ends up above the caption.
    cap = doc.add_paragraph(style="Normal")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    _insert_after(anchor, cap)

    img_p = doc.add_paragraph(style="Normal")
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(str(image_path), width=IMAGE_WIDTH)
    _insert_after(anchor, img_p)
    print(f"  inserted: {anchor_token}  <- {image_path.name}")
    return True


def main() -> None:
    doc = Document(str(MANUAL))
    inserted = 0
    for anchor_text, filename, caption in MAPPING:
        image_path = IMG_DIR / filename
        if not image_path.exists():
            print(f"  (no file yet): {filename}")
            continue
        if insert_figure(doc, anchor_text, image_path, caption):
            inserted += 1

    if inserted:
        doc.save(str(MANUAL))
        print(f"\nUpdated {MANUAL.relative_to(ROOT)} with {inserted} figure(s).")
        if ASSET_COPY.parent.exists():
            import shutil
            try:
                shutil.copyfile(MANUAL, ASSET_COPY)
                print(f"Synced {ASSET_COPY.relative_to(ROOT)}")
            except PermissionError:
                print(f"  Note: Could not overwrite {ASSET_COPY.name} because it is open in Word. Close it to sync.")
        print("Open the manual in Word and press Ctrl+A then F9 to refresh the TOC/figures.")
    else:
        print("\nNo images inserted. Drop PNGs into docs/User_manual_screenshots/ (see MAPPING).")


if __name__ == "__main__":
    main()
